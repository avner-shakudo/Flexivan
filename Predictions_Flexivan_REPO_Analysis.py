import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")  # Non-interactive backend, figures won't pop up
import matplotlib.pyplot as plt
import copy
import os
import re

from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from xgboost import XGBClassifier
from xgboost import XGBRegressor
from datetime import datetime, timedelta
from pathlib import Path
from tqdm import tqdm
from colorama import Fore, Style, init

from sklearn.model_selection import train_test_split

from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import warnings
warnings.filterwarnings("ignore")

import Flexivan_Prediction_Package


# DESCRIPTION
# This script does what Predictions_Flexivan.ipynb does but for a repository with all the code in Predictions_Flexivan.ipynb
# more encapsulated. It was used to demonstrate model capabilities and display results for the 7.11.25 meeting with Sagar.

#region GENERAL FUNCTIONS

def Divide_ARR_2_Arrays_by_Range(DF, Column_Name, Ranges):
    # This function takes a DataFrame and Ranges = [0, 10, 20, ..., 100] any selected values
    # and returns a list of arrays that contains the values between those values

    DF_LIST = []
    RANGES = []

    for ii in range(len(Ranges)-1):
        DF_LIST.append(DF[(DF[Column_Name]>=Ranges[ii]) & (DF[Column_Name]<Ranges[ii+1])])
        RANGES.append([Ranges[ii], Ranges[ii+1]])

    return DF_LIST, RANGES

def train_and_test_xgboost(df, target_column, test_size, RegressionORPrediction=0, random_state=42):
    # 1. Separate features and target
    X = df.drop(columns=[target_column])
    y = df[target_column]

    # 2. Train/test split
    Split_Index = int((1-test_size) * len(df))
    X_train = df.iloc[0:Split_Index].drop(columns=[target_column])
    y_train = df.iloc[0:Split_Index][target_column]
    X_test = df.iloc[Split_Index:].drop(columns=[target_column])
    y_test = df.iloc[Split_Index:][target_column]

    # X_train, X_test, y_train, y_test = train_test_split(
    #     X, y, test_size=test_size, random_state=random_state
    # )

    # 3. Initialize XGBoost REGRESSOR model
    if RegressionORPrediction == 0:
        model = XGBClassifier(use_label_encoder=False, eval_metric='logloss')
    else:       
        model = XGBRegressor(use_label_encoder=False, eval_metric='logloss')

    # 4. Train the model
    model.fit(X_train, y_train)

    # 5. Predict and evaluate
    y_pred = model.predict(X_test)
    # accuracy = accuracy_score(y_test, y_pred)

    # print(f"Test accuracy: {accuracy:.4f}")

    return model, y_test, y_pred
    
def Display_CDF(ARR, ax=None):
    # This function displays the CDF of a selected Numpy array - what percentage of the data is found under which value

    # Sort the data
    sorted_data = np.sort(ARR)

    # Compute CDF values
    CDF = (100 *np.arange(1, len(sorted_data) + 1) / len(sorted_data))[::-1]

    # Plot the CDF
    if ax is None:
        fig, ax = plt.subplots()
        
    ax.plot(sorted_data, CDF, linewidth=2)
    plt.grid(True, linestyle='--', alpha=0.6)
    plt.show()

    return CDF, ax
    
def enumerate_columns(df, column_name):
    """
    Replace values in each column with integer codes representing
    the unique values in that column.
    """
    unique_vals = {v: i for i, v in enumerate(df[column_name].unique())}
    df[column_name] = df[column_name].map(unique_vals)
    
    return df

def Analyze_RAW_Windows_Results(DATA, RESULTS_DETAILED_DICT):
    # This function takes the RESULTS_DETAILED_DF yielded by the sliding_xgb_window_eval function and contains the the results y_test-y_pred
    # and generates a DataFrame that contains pickup date (retrieved by iloc index from DATA) and the results of the prediction vs. GT.
    # Second part of the analysis, takes the repeating indexes (table index) and chooses the best prediction result
    # RESULTS_DETAILED_DICT = [y_test, y_pred, Original_Indexes]            (Original_Indexes are in the full DATA DF cleaned from CSV file)

    Window_Size = None
    RESULTS_DETAILED_DF = None
    DATA_COLUMNS = list(DATA.columns)

    # Collect all results from all windows
    for key in RESULTS_DETAILED_DICT.keys():
        if Window_Size is None:
            Window_Size = len(RESULTS_DETAILED_DICT[key][0])

        Col_index = DATA_COLUMNS.index('CHS Pickup Date')

        Pickup_Dates = DATA.iloc[RESULTS_DETAILED_DICT[key][2], Col_index]

        RESULTS_DETAILED_DF_TEMP = pd.DataFrame({
            'Pickup Date': np.array(Pickup_Dates).ravel(),
            'Return Date Diff': np.array(RESULTS_DETAILED_DICT[key][0]).ravel(),
            'Return Date Diff Predicted': np.array(RESULTS_DETAILED_DICT[key][1]).ravel(),
            'Diff % (ABS)': 100*abs(np.array(RESULTS_DETAILED_DICT[key][1])-np.array(RESULTS_DETAILED_DICT[key][0]))/np.array(RESULTS_DETAILED_DICT[key][0])
        })

        if RESULTS_DETAILED_DF is None:
            RESULTS_DETAILED_DF = copy.deepcopy(RESULTS_DETAILED_DF_TEMP)
        else:
            RESULTS_DETAILED_DF = pd.concat([RESULTS_DETAILED_DF, RESULTS_DETAILED_DF_TEMP])

    RESULTS_DETAILED_DF_REFINED = copy.deepcopy(RESULTS_DETAILED_DF)

    return RESULTS_DETAILED_DF_REFINED
        
def set_thick_outside_borders(table, size=12, color="000000"):
    """
    Apply thick borders to the outside of a table.
    - size: border thickness (12 ~ 2pt)
    - color: hex color code (default black)
    """
    tbl = table._tbl

    # Ensure tblPr exists
    if tbl.tblPr is None:
        tbl.tblPr = OxmlElement('w:tblPr')

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tblBorders.append(border)

    # Disable inside borders
    for inside in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{inside}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)

    # Append borders to table properties
    tbl.tblPr.append(tblBorders)

def Add_DICT_2_Table(doc, Results_DICT, HEADERS):
    # This function takes a dictionary that contains one data item in each key and adds it to a table in doc.
    # The function returns the DOCX document object with the added table

    # Add table with header row
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # You can also try 'Light Grid Accent 1', etc.

    # Add header cells
    hdr_cells = table.rows[0].cells
    for ii, header in enumerate(HEADERS):
        # hdr_cells[ii].text = header
        para = hdr_cells[ii].paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(12)
   
    for key in Results_DICT.keys():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(np.round(Results_DICT[key], 2)) + '%'

    set_thick_outside_borders(table)

    return doc

def Generate_Return_Dates_from_DIFF(Pickup_Dates, DIFFs):
    # Genetraes an array of return dates (datetime format) from datetimes in Pickup_Dates using float values in DIFFs.
    # This is used for the ACCUM file generation of the final result

    Return_Dates = []

    for ii, pickup_date in enumerate(Pickup_Dates):
        return_date = pickup_date + timedelta(days=DIFFs[ii])
        Return_Dates.append(return_date)

    return Return_Dates

def Add_File_Results_2_ACCUM_Results_File(Results_File_ACCUM_DF, Results_DF):
    if Results_File_ACCUM_DF is None:
        Results_File_ACCUM_DF = pd.DataFrame(columns=['LOT', 'Date', 'Predicted_Pickups', 'Predicted_Returns'])

    Generate_Return_Dates_from_DIFF(Results_DF[''], Results_DF[''])

    return Results_File_ACCUM_DF

#endregion

#region GENERAL PARAMETERS

# Target_Folder = '/Users/avner/flexivan/Daily prediction/DATA'
Target_Folder = '/root/Flexivan/Flexivan/Daily prediction/DATA'
Sorting_Field='CHS Pickup Date'
Error_Threshold = 20
Predicted_Column = 'Return Time Difference'
test_size = .1
Columns_2_Drop_From_Training = ['CHS ID', 'CTR Trip Id', 'CHS Return Dt', 'CHS Return LOC', 'CHS Pickup Date', 'CTR pick Dt', 'CTR Return Dt']
Test_Portion = .2
Window_Size = 5000
Window_Step_Size = int(Test_Portion * Window_Size)      # So there will be no overlap (p.24 top)
Error_THR = 20      # Percentage (0-100)

#endregion

# Getting all files for analysis
FILENAMES = [f for f in os.listdir(Target_Folder) if os.path.isfile(os.path.join(Target_Folder, f))]
selected = "Latest_Test_"
selected2 = '_DETAILED'
FILENAMES = [f for f in FILENAMES if selected in f and selected2 not in f]
DATES = Flexivan_Prediction_Package.extract_datetimes_from_filenames(FILENAMES)
Filenames_DF = pd.DataFrame({
    "Filenames": FILENAMES,
    "Dates": DATES
})
Filenames_DF_Sorted = Filenames_DF.sort_values(by='Dates')
FILENAMES = list(Filenames_DF_Sorted['Filenames'])
DATES = list(Filenames_DF_Sorted['Dates'])

#region Initializing the basic DOCX report

doc = Document()
doc.add_heading(f"Prediction Results for {Predicted_Column}\n", level=1)
p = doc.add_paragraph('This document contains the analysis of the prediction results for time difference between pickup and return of chessies. ' \
                    'The model used for the prediction is the sliding window XGBoost (parameters used are specified in the following section).')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading(f"General Analysis Parameters and Information\n", level=2)
doc.add_paragraph(f"Number of files in analysis\t{len(FILENAMES)}")
doc.add_paragraph(f"Dates range:\t\t\t{min(DATES).strftime('%Y-%m-%d')} - {max(DATES).strftime('%Y-%m-%d')}")
doc.add_paragraph(f'Columns not used:\t\t{len(Columns_2_Drop_From_Training)}')
for field in Columns_2_Drop_From_Training:
    doc.add_paragraph(f'\t\t\t{field}')
doc.add_paragraph('Model used:\t\t\tSliding window XGBoost (Regressor)')
doc.add_paragraph(f'Window size:\t\t\t{Window_Size}')
doc.add_paragraph(f'Window step size:\t\t{Window_Step_Size}')
doc.add_paragraph(f'Error THR:\t\t\t{Error_THR}')
doc.add_paragraph(f'Test portion:\t\t\t{Test_Portion}')

#endregion

#region File detailed analysis

doc.add_page_break()
doc.add_heading(f"Files Detailed Analysis\n", level=2)

# Analyzing each file and adding results to a DOCX report
Files_No = 0
print(f'Analyzing data files to predict {Predicted_Column}')

Results_Under_Error_THR = {}
FAiled_Filenames = []
Results_File_ACCUM_DF = None

for filename in tqdm(FILENAMES):
    try:
        doc.add_heading(f"Results for {filename}", level=3)

        File_Analysis_Results_OBJ = Flexivan_Prediction_Package.File_Analysis_Reults(f'{Target_Folder}/{filename}', Sorting_Field, Columns_2_Drop_From_Training)
    
        File_Analysis_Results_OBJ.Analyze_Data_File(f'{Target_Folder}/{filename}', Columns_2_Drop_From_Training, Window_Size, Window_Step_Size, Error_THR, Test_Portion, Sorting_Field)
        
        # Handle the adding of the ACCUM predictions to the main results file (daily file)
        Results_File_ACCUM_DF = Flexivan_Prediction_Package.Add_File_Results_2_ACCUM_Results_DF(File_Analysis_Results_OBJ.DATA_ORIG, Results_File_ACCUM_DF, File_Analysis_Results_OBJ.Return_DIFF_RESULTS_DETAILED_DICT)

        Results_File_ACCUM_DF = Results_File_ACCUM_DF.sort_values(by="Date")  # ascending order by default
        Results_File_ACCUM_DF.to_csv('Returns_Pickups_ACCUM_Report.csv')

        # Save prediction results to CSV
        name, ext = os.path.splitext(filename)
        print('Saving detailed results into CSV...', end='')
        File_Analysis_Results_OBJ.Return_DIFF_RESULTS_DF.to_csv(f'{Target_Folder}/{name}_DETAILED_PREDICTION_RESULTS.{ext}')
        print('DONE.')

        doc.add_paragraph(f"Lines used:\t{self.Analysis_Info_DICT['Rows_After_Cleaning']} out of {Analysis_Info_DICT['Total_Lines']} ({int(100*Analysis_Info_DICT['Rows_After_Cleaning']/Analysis_Info_DICT['Total_Lines'])}%)")
        doc.add_paragraph(f'Data sorted by:\t{Sorting_Field}')

        Results_Under_Error_THR[filename] = np.average(np.array(Results_DF['pct_under_threshold']))
        
        #region Generating results display
        
        # Create the figure with 1 row, 2 columns
        fig, axes = plt.subplots(1, 2, figsize=(12, 5))

        # --- Left plot ---
        axes[0].plot(File_Analysis_Results_OBJ.Return_DIFF_RESULTS_DF['start'], File_Analysis_Results_OBJ.Return_DIFF_RESULTS_DF['pct_under_threshold'], linewidth=2)

        # Custom x-axis labels
        X_Indexes = axes[0].get_xticks()
        X_Labels = [
            datetime.strftime(File_Analysis_Results_OBJ.DATA_ORIG.iloc[int(x)][Sorting_Field], '%Y-%m-%d') 
            if x < len(File_Analysis_Results_OBJ.DATA_ORIG) else '' 
            for x in X_Indexes
        ]
        axes[0].set_xticks(X_Indexes)
        axes[0].set_xticklabels(X_Labels, rotation=45, ha='right')

        axes[0].set_xlabel('Time', fontweight='bold', fontsize=12)
        axes[0].set_ylabel('Percentage Under Threshold', fontweight='bold', fontsize=12)
        axes[0].set_title(
            f'Sliding Window Prediction Performance\nPercentage of the Predictions Under {Error_Threshold}% Error', 
            fontweight='bold', fontsize=15
        )
        axes[0].grid()

        # --- Right plot ---
        # Assuming Display_CDF returns (CDF, ax) — pass axes[1] to it
        CDF, _ = Display_CDF(File_Analysis_Results_OBJ.Return_DIFF_RESULTS_DF['pct_under_threshold'].sort_values(), axes[1])

        axes[1].set_xlabel(f'Percentage Under {Error_Threshold}% \nError Threshold[%]', fontweight='bold', fontsize=12)
        axes[1].set_ylabel('Percentage [%]', fontweight='bold', fontsize=12)
        axes[1].set_title('Sliding Window Prediction Performance CDF', fontweight='bold', fontsize=15)

        # Adjust layout to avoid overlapping labels
        plt.tight_layout()

        # Save the figure to a file WITHOUT displaying it
        img_path = f'{Target_Folder}/sales_chart.png'
        fig.savefig(img_path, bbox_inches='tight', dpi=300)
        doc.add_picture(img_path, width=Inches(6))  # adjust width as needed

        # Close the figure to free memory
        plt.close(fig)

        #endregion
    except Exception as e:
        print(Fore.RED + f'Exception: {e}' + Fore.RESET)
        FAiled_Filenames.append(filename)

    Files_No += 1

    # if Files_No>=2:
    #     break

try: 
    os.remove(f'{Target_Folder}/sales_chart.png')
except Exception as e:
    print(f'ERROR: {e}')

#endregion

#region Adding the results summary appendix

doc.add_page_break()
doc.add_heading(f"Results Summary\n", level=2)
p = doc.add_paragraph(f'This section contains a results summary table for all files. The summary metric is the percentage of the prediction results that were below \
                      the selected error threshold ({Error_THR}%) averaged over all windows in file analysis')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
HEADERS = ['Data Filename', f'% Below {Error_THR}%']
doc = Add_DICT_2_Table(doc, Results_Under_Error_THR, HEADERS)

Results_Under_Error_THR

#endregion

# Finalizing document
print('Finalizing report...', end='')
doc.save(f"{Target_Folder}/Flexivan_Shakudo_Results.docx")
print('DONE.')

if len(FAiled_Filenames)>0:
    print('the following files failed analysis:')

    for ii, filename in enumerate(FAiled_Filenames):
        print(f'File No.{ii+1}:\t{filename}')

print('\nANALYSIS COMPLETE\n')

#endregion
