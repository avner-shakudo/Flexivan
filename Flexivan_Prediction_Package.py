import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")  # Non-interactive backend, figures won't pop up
import matplotlib.pyplot as plt
import copy
import os
import re
import sys

from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from sklearn.preprocessing import LabelEncoder

from xgboost import XGBClassifier
from xgboost import XGBRegressor
import xgboost as xgb
from datetime import datetime, timedelta
from pathlib import Path
from tqdm import tqdm
from colorama import Fore, Style, init

from sklearn.model_selection import train_test_split

from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import warnings
warnings.filterwarnings("ignore")

class File_Analysis_Reults():
    def __init__(self, Fullpath, Sorting_Field, Columns_2_Drop_From_Training, Enumerated_Columns_LIST=None, Units='Days'):
        self.Fullpath = Fullpath
        self.Sorting_Field = Sorting_Field
        self.Columns_2_Drop_From_Training = Columns_2_Drop_From_Training
        self.Enumerated_Columns_LIST = Enumerated_Columns_LIST
        self.Units = Units
        self.Folder = None
        self.Filename = None
        self.Extension = None
        self.DATA_ORIG = None               # The raw DataFrame that contains the data for analysis (loaded from CSV file)
        self.DATA = None                    # Contains the same data after cleaning (dropping columns and null rows etc.)
        self.Analysis_Info_DICT = {}        # A DICT contains general info about the data and results for use outside the function 
        self.Return_DIFF_y_test = None
        self.Return_DIFF_y_pred = None
        self.Return_DIFF_RESULTS_DETAILED_DICT = None
        self.Return_DIFF_RESULTS_DF = None
        self.LOT_Prediction_y_test = None
        self.LOT_Prediction_y_pred = None
        self.LOT_Prediction_RESULTS_DETAILED_DICT = None
        self.MAPPINGS_CHS_Pickup_Loc = {}

        self.Load_Data(Fullpath)
        self.Clean_Data(Sorting_Field)
        # self.Enumerate_Data(self.Enumerated_Columns_LIST)
        
    def Load_Data(self, Fullpath):
        if isinstance(Fullpath, str):
            Fullpath = [Fullpath]

        self.DATA_ORIG = None

        for fullpath in Fullpath:
            DATA_TEMP = pd.read_csv(fullpath)

            if self.DATA_ORIG is None:
                self.DATA_ORIG = copy.deepcopy(DATA_TEMP)
            else:
                self.DATA_ORIG = pd.concat([self.DATA_ORIG, DATA_TEMP])

        self.DATA = copy.deepcopy(self.DATA_ORIG)

        if isinstance(Fullpath, list):
            self.Folder = None
            self.Filename = None
            self.Extension = None
        else:
            self.Folder = os.path.dirname(Fullpath)
            self.Filename = os.path.splitext(os.path.basename(Fullpath))[0]
            self.Extension = os.path.splitext(Fullpath)[1]
        
        return 0

    def Clean_Data(self, Sorting_Field):
        # print(f'Cleaning data and sorting by {Sorting_Field}')
        self.DATA['CHS Pickup Date'] = pd.to_datetime(self.DATA['CHS Pickup Date'], errors='coerce')
        self.DATA['CHS Return Dt'] = pd.to_datetime(self.DATA['CHS Return Dt'], errors='coerce')

        for col in self.DATA.select_dtypes(include=['object']).columns:
            self.DATA[col] = self.DATA[col].map(lambda x: x.strip() if isinstance(x, str) else x)

        # Normalize common missing markers and empty strings to NaN
        self.DATA.replace(['', 'NA', 'N/A', 'na', 'n/a'], np.nan, inplace=True)

        # Drop any row that contains at least one NaN
        self.DATA.dropna(axis=0, how='any', inplace=True)

        # print(f'Sorting by field {Sorting_Field}...', end='')
        self.DATA = self.DATA.sort_values(by=Sorting_Field)

        self.DATA.reset_index(drop=True, inplace=True)

        return 0

    def Enumerate_Data(self, Enumerated_Columns_LIST=None):
       # If no list was provided, use your default fields:
        if Enumerated_Columns_LIST is None:
            Enumerated_Columns_LIST = [
                'CHS Pickup Loc', 'CHS Return Loc', 'CHS pickup MCO', 
                'CTR Trip MCO', 'O Customer', 'Customer', 'DC Loc', 
                'CTR Pickup Term', 'CTR Return Term', 'pgkey', 
                'CTR Trip Loc Type Pattern', 'CTR Trip Pattern'
            ]
            
        # Convert datetime columns to timestamps (seconds)
        for col in self.DATA.select_dtypes(include=['datetime64[ns]', 'datetime64']).columns:
            self.DATA[col] = self.DATA[col].astype('int64') // 1_000_000_000

        # Enumerate specified columns
        for col in Enumerated_Columns_LIST:
            if col in self.DATA.columns:
                self.DATA[col] = self.DATA[col].astype('category')
                # Save mapping
                if col=='CHS Pickup Loc':
                    self.MAPPINGS_CHS_Pickup_Loc = dict(enumerate(self.DATA[col].cat.categories))
                # Convert to codes
                self.DATA[col] = self.DATA[col].cat.codes

        # Enumerate remaining object columns
        for col in self.DATA.select_dtypes(include=['object']).columns:
            self.DATA[col] = self.DATA[col].astype('category')
            # Convert to codes
            self.DATA[col] = self.DATA[col].cat.codes

        return 0

    def Analyze_Data_File(self, window_size, step, Error_Threshold=20, test_frac=.2):
        # This function analyzes data in CSV found in Fullpath. It uses the sliding window XGBoost model (test and train over a portion of the data)
        # The process is repeated for sliding windows of size (window_size with step of step samples)

        Folder = os.path.dirname(self.Fullpath)
        Filename = os.path.basename(self.Fullpath)
        Analysis_Info_DICT = {}

        self.Clean_Data(self.Sorting_Field)
        print(f"Rows after cleaning: " + Fore.YELLOW + f'{len(self.DATA)}' + Fore.RESET + ' (' + Fore.GREEN + f'{int(100*len(self.DATA)/len(self.DATA_ORIG))}' + Fore.RESET + ')% remained after cleaning')
        print(f'Sorted by ' + Fore.YELLOW + f'{self.Sorting_Field}' + Fore.RESET + ' field')

        self.Analysis_Info_DICT['Total_Lines'] = len(self.DATA_ORIG)
        self.Analysis_Info_DICT['Rows_After_Cleaning'] = len(self.DATA)
    
        #region RETURN PREDICTIONS

        #region Calculating differences and addting new columns
        
        Diff_Col_Name = f'Pickup_Return_Time_Diff_{self.Units}'
        Analysis_Info_DICT['Time_Diff_Units'] = self.Units

        if self.Units == 'Hours':
            self.DATA[Diff_Col_Name] = (self.DATA['CHS Return Dt'] - self.DATA['CHS Pickup Date']).dt.total_seconds() / 3600
            
        elif self.Units == 'Days':
            self.DATA[Diff_Col_Name] = (self.DATA['CHS Return Dt'] - self.DATA['CHS Pickup Date']).dt.total_seconds() / (3600*24)

        #endregion

        self.Enumerate_Data()

        # Predicting return time diff
        step=int(test_frac * window_size)      # Force so there will be no overlapping results (more than one prediction per sample)
        self.Return_DIFF_RESULTS_DF, self.Return_DIFF_y_test, self.Return_DIFF_y_pred, self.Return_DIFF_RESULTS_DETAILED_DICT = sliding_xgb_window_eval(self.DATA, Diff_Col_Name, window_size, step, test_frac,
                                                                                Error_Threshold, xgb_params=None, random_state=42,
                                                                                min_test_samples=2, show_progress=False, Classifier_or_Regressor=0)
        #endregion
        
        #region LOR PREDICTION

        # # Predicting return LOT
        # __, y_test_Return_LOT, y_pred_Return_LOT, self.RESULTS_DETAILED_DICT_LOT = sliding_xgb_window_eval(self.DATA, 'CHS Return Loc', window_size, step, test_frac,
        #                                                                         Error_Threshold, xgb_params=None, random_state=42,
        #                                                                         min_test_samples=2, show_progress=False, Classifier_or_Regressor=0)
        # y_pred_Return_LOT = [int(x) for x in y_pred_Return_LOT]

        #endregion

        # # Analyze RESULTS_DETAILE
        self.Analysis_Info_DICT['Model_Window_Size'] = window_size
        self.Analysis_Info_DICT['Model_Window_Step'] = step
        self.Analysis_Info_DICT['Model_Erro_THR'] = Error_Threshold

        return 0
 

def Divide_ARR_2_Arrays_by_Range(DF, Column_Name, Ranges):
    # This function takes a DataFrame and Ranges = [0, 10, 20, ..., 100] any selected values
    # and returns a list of arrays that contains the values between those values

    DF_LIST = []
    RANGES = []

    for ii in range(len(Ranges)-1):
        DF_LIST.append(DF[(DF[Column_Name]>=Ranges[ii]) & (DF[Column_Name]<Ranges[ii+1])])
        RANGES.append([Ranges[ii], Ranges[ii+1]])

    return DF_LIST, RANGES

def train_and_test_xgboost(df, target_column, test_size, RegressionORPrediction=0, random_state=42):
    # 1. Separate features and target
    X = df.drop(columns=[target_column])
    y = df[target_column]

    # 2. Train/test split
    Split_Index = int((1-test_size) * len(df))
    X_train = df.iloc[0:Split_Index].drop(columns=[target_column])
    y_train = df.iloc[0:Split_Index][target_column]
    X_test = df.iloc[Split_Index:].drop(columns=[target_column])
    y_test = df.iloc[Split_Index:][target_column]

    # X_train, X_test, y_train, y_test = train_test_split(
    #     X, y, test_size=test_size, random_state=random_state
    # )

    # 3. Initialize XGBoost REGRESSOR model
    if RegressionORPrediction == 0:
        model = XGBClassifier(use_label_encoder=False, eval_metric='logloss')
    else:       
        model = XGBRegressor(use_label_encoder=False, eval_metric='logloss')

    # 4. Train the model
    model.fit(X_train, y_train)

    # 5. Predict and evaluate
    y_pred = model.predict(X_test)
    # accuracy = accuracy_score(y_test, y_pred)

    # print(f"Test accuracy: {accuracy:.4f}")

    return model, y_test, y_pred
    
def Display_CDF(ARR, ax=None):
    # This function displays the CDF of a selected Numpy array - what percentage of the data is found under which value

    # Sort the data
    sorted_data = np.sort(ARR)

    # Compute CDF values
    CDF = (100 *np.arange(1, len(sorted_data) + 1) / len(sorted_data))[::-1]

    # Plot the CDF
    if ax is None:
        fig, ax = plt.subplots()
        
    ax.plot(sorted_data, CDF, linewidth=2)
    plt.grid(True, linestyle='--', alpha=0.6)
    plt.show()

    return CDF, ax
    
def enumerate_columns(df, column_name):
    """
    Replace values in each column with integer codes representing
    the unique values in that column.
    """
    unique_vals = {v: i for i, v in enumerate(df[column_name].unique())}
    df[column_name] = df[column_name].map(unique_vals)
    
    return df

def sliding_xgb_window_eval(df, target_col, window_size, step, test_frac=0.2,
                            error_threshold=20.0, xgb_params=None, random_state=42,
                            min_test_samples=2, show_progress=False, Classifier_or_Regressor=0):
    """
    Slide a fixed-size window over df, train XGBRegressor on the first (1-test_frac)
    fraction and test on the last test_frac fraction. For each window compute the
    percentage of test rows whose percentage error <= error_threshold.

    Returns a DataFrame with columns:
    window_idx, start, end, n_test, pct_under_threshold
    """
    import math
    from tqdm import tqdm
    from xgboost import XGBRegressor, XGBClassifier

    if xgb_params is None:
        xgb_params = {'n_estimators': 100, 'random_state': random_state, 'verbosity': 0}

    results = []
    n = len(df)
    if window_size > n:
        window_size = n
        print("window_size larger than dataframe length. Setting one window")

    indices = range(0, n - window_size + 1, step)
    iterator = tqdm(indices) if show_progress else indices
    wi = 0
    RESULTS_DETAILED_DF = {}            # Holds the detailed results for each window for results CSV purposes. Keys are the start index,
                                        # and the content is a list of [[y_test], [y_pred]]

    print('Training and predicting for sliding windows...')
    for start in tqdm(iterator):
        y_test = None
        y_pred = None 
        pct_under = 0

        try:
            end = start + window_size  # exclusive
            window = df.iloc[start:end].copy()

            # drop rows missing target
            window = window.dropna(subset=[target_col])
            if len(window) < 2:
                continue

            # prepare features (dummies on full window to keep alignment)
            X_all = pd.get_dummies(window.drop(columns=[target_col]), drop_first=True)
            y_all = window[target_col].astype(float).values
            # y_all = window[target_col].astype(int).values

            split_idx = int(math.floor((1.0 - test_frac) * len(window)))
            # ensure at least min_test_samples in test
            if len(window) - split_idx < min_test_samples or split_idx < 1:
                # skip window too small
                continue

            X_train = X_all.iloc[:split_idx, :].values
            y_train = y_all[:split_idx]
            X_test = X_all.iloc[split_idx:, :].values
            y_test = y_all[split_idx:]

            # train
            if Classifier_or_Regressor:
                xgb_params = {'n_estimators': 100, 'random_state': random_state, 'verbosity': 0, 'num_class': len(set(y_train))}
                model = XGBClassifier(**xgb_params)
                y_test = [int(x) for x in y_test]
                y_train = [int(x) for x in y_train]
            else:
                model = XGBRegressor(**xgb_params)

            # try:
            #     model.fit(X_train, y_train)   
            # except Exception as e:
            #     pass
            model.fit(X_train, np.array(y_train))
            y_pred = model.predict(X_test)

            # model, y_test, y_pred = train_and_test_xgboost(window, target_col, test_size, RegressionORPrediction=1, random_state=42)

            # compute percent error safely:
            abs_diff = np.abs(y_test - y_pred)
            # If true value is zero, define percent error as 0 if prediction equals, else 100.
            pct_err = np.where(
                np.isclose(y_test, 0.0),
                np.where(np.isclose(abs_diff, 0.0), 0.0, 100.0),
                100.0 * abs_diff / np.abs(y_test)
            )

            RESULTS_DETAILED_DF[start] = [y_test, y_pred, start + np.arange(split_idx, split_idx+len(y_test))]

            pct_under = 100.0 * float(np.sum(pct_err <= error_threshold)) / len(pct_err)
        except:
            pass

        results.append({
            'window_idx': wi,
            'start': start,
            'end': end,
            'n_test': len(pct_err),
            'pct_under_threshold': pct_under
        })

        wi += 1

    print('DONE.')

    try:
        DF = pd.DataFrame(results).sort_values('window_idx').reset_index(drop=True), y_test, y_pred, RESULTS_DETAILED_DF
    except:
        DF = None, None, None, None

    return DF

def Analyze_RAW_Windows_Results(DATA, RESULTS_DETAILED_DICT):
    # This function takes the RESULTS_DETAILED_DF yielded by the sliding_xgb_window_eval function and contains the the results y_test-y_pred
    # and generates a DataFrame that contains pickup date (retrieved by iloc index from DATA) and the results of the prediction vs. GT.
    # Second part of the analysis, takes the repeating indexes (table index) and chooses the best prediction result
    # RESULTS_DETAILED_DICT = [y_test, y_pred, Original_Indexes]            (Original_Indexes are in the full DATA DF cleaned from CSV file)

    Window_Size = None
    RESULTS_DETAILED_DF = None
    DATA_COLUMNS = list(DATA.columns)

    # Collect all results from all windows
    for key in RESULTS_DETAILED_DICT.keys():
        if Window_Size is None:
            Window_Size = len(RESULTS_DETAILED_DICT[key][0])

        Col_index = DATA_COLUMNS.index('CHS Pickup Date')

        Pickup_Dates = DATA.iloc[RESULTS_DETAILED_DICT[key][2], Col_index]

        RESULTS_DETAILED_DF_TEMP = pd.DataFrame({
            'Pickup Date': np.array(Pickup_Dates).ravel(),
            'Return Date Diff': np.array(RESULTS_DETAILED_DICT[key][0]).ravel(),
            'Return Date Diff Predicted': np.array(RESULTS_DETAILED_DICT[key][1]).ravel(),
            'Diff % (ABS)': 100*abs(np.array(RESULTS_DETAILED_DICT[key][1])-np.array(RESULTS_DETAILED_DICT[key][0]))/np.array(RESULTS_DETAILED_DICT[key][0])
        })

        if RESULTS_DETAILED_DF is None:
            RESULTS_DETAILED_DF = copy.deepcopy(RESULTS_DETAILED_DF_TEMP)
        else:
            RESULTS_DETAILED_DF = pd.concat([RESULTS_DETAILED_DF, RESULTS_DETAILED_DF_TEMP])

    RESULTS_DETAILED_DF_REFINED = copy.deepcopy(RESULTS_DETAILED_DF)

    return RESULTS_DETAILED_DF_REFINED
        
def extract_datetimes_from_filenames(filenames):
    """
    Extracts datetime objects from filenames of the form 'Latest_Test_<DATE>'.
    Supports formats like YYYY-MM-DD, YYYYMMDD, or YYYY_MM_DD.
    """
    datetimes = []

    for name in filenames:
        # Try to find the date pattern
        match = re.search(r'(\d{4}[-_]\d{2}[-_]\d{2}|\d{8})', name)
        if match:
            date_str = match.group(1)
            
            # Normalize formats like YYYY_MM_DD → YYYY-MM-DD
            date_str = date_str.replace("_", "-")
            
            # Try multiple possible formats
            for fmt in ("%Y-%m-%d", "%Y%m%d"):
                try:
                    dt = datetime.strptime(date_str, fmt)
                    datetimes.append(dt)
                    break
                except ValueError:
                    continue

    return datetimes

def set_thick_outside_borders(table, size=12, color="000000"):
    """
    Apply thick borders to the outside of a table.
    - size: border thickness (12 ~ 2pt)
    - color: hex color code (default black)
    """
    tbl = table._tbl

    # Ensure tblPr exists
    if tbl.tblPr is None:
        tbl.tblPr = OxmlElement('w:tblPr')

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tblBorders.append(border)

    # Disable inside borders
    for inside in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{inside}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)

    # Append borders to table properties
    tbl.tblPr.append(tblBorders)

def Add_DICT_2_Table(doc, Results_DICT, HEADERS):
    # This function takes a dictionary that contains one data item in each key and adds it to a table in doc.
    # The function returns the DOCX document object with the added table

    # Add table with header row
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'  # You can also try 'Light Grid Accent 1', etc.

    # Add header cells
    hdr_cells = table.rows[0].cells
    for ii, header in enumerate(HEADERS):
        # hdr_cells[ii].text = header
        para = hdr_cells[ii].paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(12)
   
    for key in Results_DICT.keys():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(np.round(Results_DICT[key], 2)) + '%'

    set_thick_outside_borders(table)

    return doc

def Generate_Return_Dates_from_DIFF(Pickup_Dates, DIFFs):
    # Genetraes an array of return dates (datetime format) from datetimes in Pickup_Dates using float values in DIFFs.
    # This is used for the ACCUM file generation of the final result

    Return_Dates = []

    for ii, pickup_date in enumerate(Pickup_Dates):
        return_date = pickup_date + timedelta(days=DIFFs[ii])
        Return_Dates.append(return_date)

    return Return_Dates

def Add_File_Results_2_ACCUM_Results_DF(DATA_ORIG, Results_File_ACCUM_DF, Detailed_Pred_Results_DF):
    # DATA_ORIG - the original cleaned version of the data file. This is used to get information about a sample that
    #             does not exist in the Detailed_Pred_Results_DF table (LOT and PU date)
    # Results_File_ACCUM_DF - this is the daily report to maintain with ACCUM figure for returns and pickups
    # Detailed_Pred_Results_DF = [[y_pred], [y_test]], [Indexes_in_ORIG], [Predicted_LOT]]

    if Results_File_ACCUM_DF is None or len(Results_File_ACCUM_DF)==0:
        Results_File_ACCUM_DF = pd.DataFrame(columns=['Date', 'Predicted_Pickups', 'Predicted_Returns'])
    
    #region Adding the RETURNS PREDICTION indormation into the ACCUM report

    for jj, key in enumerate(Detailed_Pred_Results_DF.keys()):
        Epoch_DATA = Detailed_Pred_Results_DF[key]      # [[y_pred], [y_test]], [Indexes_in_ORIG]]

        y_test = Epoch_DATA[0]
        y_pred = Epoch_DATA[1]
        Indexes_ORIG = Epoch_DATA[2]
        LOT = 0             # TODO: complete the LOT prediction model
        # LOTS = Epoch_DATA[3]
                                    # These are all aligned in indexes
    
        PickUp_Dates = DATA_ORIG.iloc[Indexes_ORIG]['CHS Pickup Date']
        Display_Factor = 100

        for ii, DATE in enumerate(PickUp_Dates):
            if ii%Display_Factor==0:
                print(f'Epoch No.{jj+1} out of {len(Detailed_Pred_Results_DF.keys())} - {ii} out of {len(PickUp_Dates)} Dates completed', end='\r')

            # LOT = LOTS[ii]
            try:
                DATE = datetime.strptime(DATE, "%Y-%m-%d %H:%M:%S")
            except:
                pass

            Future_Date = DATE + timedelta(days=float(y_pred[ii]))
            Future_Date = DATE.date()
            Future_Date = DATE.strftime("%Y-%m-%d")

            if len(Results_File_ACCUM_DF)==0:
                Results_File_ACCUM_DF.loc[len(Results_File_ACCUM_DF)] = [Future_Date, 0, 1]
                Results_File_ACCUM_DF.set_index('Date', inplace=True)
            else:
                try:
                    Results_File_ACCUM_DF.loc[Future_Date, 'Predicted_Returns'] += 1
                    pass
                except:
                    Results_File_ACCUM_DF.loc[Future_Date] = [0, 1]
    #endregion

    return Results_File_ACCUM_DF

def extract_datetimes_from_filenames(filenames):
    """
    Extracts datetime objects from filenames of the form 'Latest_Test_<DATE>'.
    Supports formats like YYYY-MM-DD, YYYYMMDD, or YYYY_MM_DD.
    """
    datetimes = []

    for name in filenames:
        # Try to find the date pattern
        match = re.search(r'(\d{4}[-_]\d{2}[-_]\d{2}|\d{8})', name)
        if match:
            date_str = match.group(1)
            
            # Normalize formats like YYYY_MM_DD → YYYY-MM-DD
            date_str = date_str.replace("_", "-")
            
            # Try multiple possible formats
            for fmt in ("%Y-%m-%d", "%Y%m%d"):
                try:
                    dt = datetime.strptime(date_str, fmt)
                    datetimes.append(dt)
                    break
                except ValueError:
                    continue

    return datetimes

def Comma_Separation_Num_String(n):
    return f"{n:,}"

def align_df_to_model(df: pd.DataFrame, model, fill_value=0):
    """
    Align a dataframe to match the feature order and feature names
    expected by an XGBoost model (XGBClassifier or Booster).
    """

    # Extract the booster
    if hasattr(model, "get_booster"):
        booster = model.get_booster()
    else:
        booster = model  # assume it's already a booster

    expected_cols = booster.feature_names

    if expected_cols is None:
        raise ValueError(
            "Model has no feature names. "
            "Train the model with a DataFrame or specify feature_names in DMatrix."
        )

    # Add any missing columns
    for col in expected_cols:
        if col not in df.columns:
            df[col] = fill_value

    # Keep only expected columns in correct order
    df = df[expected_cols]

    return df

def Supplument_XGBoost_Model_TRAINing(Booster_Model, X_new, y_new):
    training_cols = Booster_Model.feature_names
    X_new = align_to_train_columns(X_new, training_cols)
    dnew = xgb.DMatrix(X_new, label=y_new)

    cols = Booster_Model.feature_names

    params = {
        "objective": "multi:softprob",   # or "binary:logistic" for binary
        "num_class": len(set(y_new)),        # remove this line if binary
    }

    # Continue training for some rounds
    Booster_Model = xgb.train(
        params,
        dnew,
        num_boost_round=20,       # number of NEW trees
        xgb_model=Booster_Model         # important!
    )

    return Booster_Model

def Train_XGBoost(X, y):
    le = LabelEncoder()
    y_encoded = le.fit_transform(y)

    dtrain = xgb.DMatrix(X, label=y_encoded)

    params = {
        "objective": "multi:softprob",   # or "binary:logistic" for binary
        "num_class": len(set(y)),        # remove this line if binary
    }

    model = xgb.train(
        params=params,
        dtrain=dtrain,
        num_boost_round=20
    )

    return model, le

def Predict_XGBoost(model, LE, X_test):
    train_cols = model.feature_names
    X_test = align_to_train_columns(X_test, train_cols)
    dtest = xgb.DMatrix(X_test)
    y_pred_probs = model.predict(dtest)

    # binary vs multi-class
    if y_pred_probs.ndim == 1:
        single_class = LE.classes_[0]
        y_pred_labels = np.array([single_class] * len(X_test))
    else:
        # multi-class case
        y_pred_class = np.argmax(y_pred_probs, axis=1)
        
        # convert back to original labels
        y_pred_labels = LE.inverse_transform(y_pred_class)

    return y_pred_labels


def align_to_train_columns(X_new, train_cols):
    X_new = X_new.copy()
    
    # 1. Add missing columns
    for col in train_cols:
        if col not in X_new.columns:
            X_new[col] = 0  # or np.nan
    
    # 2. Remove extra columns
    X_new = X_new[train_cols]
    
    return X_new

def split_df_by_membership(df, column_name, short_vec):
    # If short_vec is empty, return empty df_in and full df_out
    if len(short_vec) == 0:
        return df.iloc[0:0].copy(), df.copy()
    
    # Convert short_vec to a set for fast lookup
    allowed = set(short_vec)
    
    df_in = df[df[column_name].isin(allowed)].copy()
    df_out = df[~df[column_name].isin(allowed)].copy()
    
    return df_in, df_out


def Detect_File_Index_4_Stable_Enum(FILENAMES, Field, Folder=None):
    # This function finds the file index in FILENAMES that is the last file that new values of Field were seen.
    # From that file onwards, all files in FILENAMES will contain the subset of the Field unique values that were
    # accumulated up to that point.

    File_Index = 0
    Unique_Values = []

    print(f'Finding last file that contains stable unique values of {Field}...')

    for filename in tqdm(FILENAMES):
        if Folder is None:
            DATA = pd.read_csv(filename)
        else:
            DATA = pd.read_csv(f'{Folder}/{filename}')
            DATA_Unique_Values = list(DATA[Field].unique())

            combined = list(set(Unique_Values + DATA_Unique_Values))

            if len(combined) > len(Unique_Values):
                File_Index += 1
                Unique_Values = copy.deepcopy(combined)

    print('DONE')

    return File_Index

def map_column_inplace(df, column_name, mapping_dict, fill_value=-1):
    """
    Replace a DataFrame column with integers according to a provided mapping dictionary.
    
    Args:
        df: pandas DataFrame
        column_name: str, name of the column to map
        mapping_dict: dict, {string_value: int_value}
        fill_value: int, value to assign if string not in mapping_dict (default -1)
    
    Returns:
        pandas DataFrame with the column overwritten by mapped integers
    """
    df[column_name] = df[column_name].map(mapping_dict).fillna(fill_value).astype(int)
    
    return df

def extrapolate_value_auto_degree(x, y, x_target, max_degree=1):
    x = np.array(x)
    y = np.array(y)
    
    if len(x) == 0:
        return -1, -1
    elif len(x) == 1:
        # Only one point: return that Y value
        return y[0], 0
    
    best_degree = 1
    best_error = float('inf')
    best_poly = None
    
    # Limit degree to number of points minus 1
    max_degree = min(max_degree, len(x)-1)
    
    for degree in range(1, max_degree+1):
        coeffs = np.polyfit(x, y, degree)
        poly = np.poly1d(coeffs)
        y_fit = poly(x)
        error = np.mean((y - y_fit)**2)
        
        if error < best_error:
            best_error = error
            best_degree = degree
            best_poly = poly

    # Ensure best_poly is assigned
    if best_poly is None:
        # fallback to linear
        coeffs = np.polyfit(x, y, 1)
        best_poly = np.poly1d(coeffs)
        best_degree = 1

    y_target = best_poly(x_target)
    return y_target, best_degree


def find_latest_date(date_dict):
    """
    Return the latest datetime from dictionary keys.
    Keys can be datetime objects, date objects, or date/datetime strings.
    Supports multiple common date/datetime formats.
    """
    if not date_dict:
        return None

    # Known date/datetime formats
    date_formats = [
        "%Y-%m-%d",
        "%Y/%m/%d",
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%Y-%m-%d %H:%M:%S",
        "%Y/%m/%d %H:%M:%S",
        "%Y-%m-%dT%H:%M:%S",  # ISO format
    ]

    parsed_datetimes = []

    for key in date_dict.keys():
        if isinstance(key, datetime):
            parsed_datetimes.append(key)
        elif isinstance(key, date):
            # Convert date to datetime at midnight
            parsed_datetimes.append(datetime.combine(key, datetime.min.time()))
        elif isinstance(key, str):
            parsed = None
            for fmt in date_formats:
                try:
                    parsed = datetime.strptime(key, fmt)
                    break
                except ValueError:
                    continue
            if parsed is None:
                raise ValueError(f"Unsupported date format: {key}")
            parsed_datetimes.append(parsed)
        else:
            raise TypeError(f"Unsupported key type: {type(key)}")

    return max(parsed_datetimes)


def Extrapolate_PUs_Number(LOT_PUs_Num_Comb, Selected_Date_4_Pred, Extrapolation_Window_Size=10):
    # This function takes a LOT code and its comb of values for the number of PUs for the last N days
    # and extrapolates the number of PUs in the Selected_Date_4_Pred
    # LOT_PUs_Num_Comb[date] = [a1, a2,...,aN]
    
    try:
        values = [value for key, value in sorted(LOT_PUs_Num_Comb.items(), key=lambda x: x[0])]
        # Compute the number of time units (days) between the Selected_Date_4_Pred and the last date in COMB
        # steps_ahead = int((Selected_Date_4_Pred - find_latest_date(LOT_PUs_Num_Comb)).total_seconds() / (24*3600))

        if len(values) >= Extrapolation_Window_Size:
            # Extrapolate according to COMB figures
            return auto_extrapolate(values, 1)
        else:
            return sum(values) / len(values)
    except:
        pass

    return None

def Update_PUs_Time_Line(PU_Date_Time_Line, DATA, PU_Date_Field_Name='CHS Pickup Date'):
    # This function updates the time line of date baskets that contain how many PUs were done for each date,
    # This is done because a date file of raw samples, contains different dates and all not all PUs in the file
    # were done on that day

    try:
        DATA[PU_Date_Field_Name] = pd.to_datetime(DATA[PU_Date_Field_Name])
    except:
        pass

    DATA[PU_Date_Field_Name] = DATA[PU_Date_Field_Name].dt.date
    DATA[PU_Date_Field_Name] = pd.to_datetime(DATA[PU_Date_Field_Name])

    DATES_Unique = list(set(DATA[PU_Date_Field_Name]))

    for DATE in DATES_Unique:
        DATA_TEMP = DATA[DATA[PU_Date_Field_Name]==DATE]

        try:
            PU_Date_Time_Line[DATE] += len(DATA_TEMP)
        except:
            PU_Date_Time_Line[DATE] = len(DATA_TEMP)
        
    return PU_Date_Time_Line

def Get_Dates_COMB_From_Dates_Timeline(PU_Date_Time_Line, target_datetime, Comb_Size):
    """
    Returns a dictionary of the last N datetime-value pairs before target_datetime,
    with keys as datetime objects, sorted descending by datetime.

    Parameters:
    - PU_Date_Time_Line: dict with datetime keys and values
    - target_datetime: datetime object
    - Comb_Size: int, number of items to return
    """
    # Filter items strictly before target_datetime
    filtered_items = [
        (k, v) for k, v in PU_Date_Time_Line.items() if k < target_datetime
    ]

    # Sort by datetime descending (most recent first)
    filtered_items.sort(key=lambda x: x[0], reverse=True)

    # Take the first Comb_Size items
    selected_items = filtered_items[:Comb_Size]

    # Return as a dictionary
    return dict(selected_items)

def distance_to_closest(date, date_list):
    """
    Return the distance in days between a given date and the closest date in a list.

    Parameters:
        date (datetime): the date to compare
        date_list (list of datetime): list of datetime objects to compare against

    Returns:
        int: distance in days to the closest date
    """
    if not date_list:
        return None  # no dates to compare
    
    # Compute the absolute difference in days to each date
    distances = [abs((date - d).days) for d in date_list]
    
    # Return the minimum distance
    return min(distances)

from datetime import datetime

def map_dates_to_floats(dates):
    """
    Map a list of datetime objects to floats where:
    - earliest date is 0
    - each day difference is 1
    
    Parameters:
        dates (list of datetime): list of datetime objects
    
    Returns:
        list of float: mapped values
    """
    if not dates:
        return []
    
    # Find the earliest date
    min_date = min(dates)
    
    # Compute the difference in days as floats
    mapped_values = [(d - min_date).days for d in dates]
    
    return mapped_values

from datetime import datetime

def get_previous_N_entries(data_dict, N, target_date):
    """
    Get the 10 dictionary entries with keys preceding target_date.

    Parameters:
        data_dict (dict): keys are datetime, values are any
        target_date (datetime): reference date

    Returns:
        dict: sorted dictionary with up to 10 entries before target_date
    """
    # Filter keys that are before the target date
    preceding_items = {k: v for k, v in data_dict.items() if k < target_date}
    
    # Sort by datetime ascending
    sorted_items = dict(sorted(preceding_items.items(), key=lambda item: item[0]))
    
    # Take the last 10 items (closest 10 dates before target)
    last_N_items = dict(list(sorted_items.items())[-N:])
    
    return last_N_items
